import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_report(df, filename="trade_report.docx", include_ai=True):
    """
    导出交易报告到 Word (v7.0 Pro)
    
    :param df: 交易数据 DataFrame
    :param filename: 保存的文件名
    :param include_ai: 是否包含 AI 点评 (False = 原始数据模式)
    """
    doc = Document()
    
    # === 1. 文档标题 ===
    heading = doc.add_heading('交易复盘深度报告', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加导出时间
    from datetime import datetime
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'报告类型: {"完整复盘 (含AI审计)" if include_ai else "原始数据档案 (无干扰)"}')
    doc.add_paragraph('---')
    
    # === 2. 统计摘要 ===
    total_trades = len(df)
    win_trades = len(df[df['net_pnl'] > 0])
    total_pnl = df['net_pnl'].sum()
    win_rate = (win_trades / total_trades * 100) if total_trades > 0 else 0
    
    stats_para = doc.add_paragraph()
    stats_para.add_run(f'总交易笔数: {total_trades} | ').bold = True
    stats_para.add_run(f'总盈亏: ${total_pnl:.2f} | ').bold = True
    stats_para.add_run(f'胜率: {win_rate:.1f}%').bold = True
    
    # === 3. 逐笔交易详情 ===
    # 按平仓时间倒序排列
    if 'close_time' in df.columns:
        df = df.sort_values(by='close_time', ascending=False)
        
    for index, row in df.iterrows():
        # 分隔符
        doc.add_paragraph('_' * 40)
        
        # 交易标题 (Symbol + Direction + PnL)
        pnl = row.get('net_pnl', 0)
        symbol = row.get('symbol', 'Unknown')
        direction = row.get('direction', 'N/A')
        date_str = row.get('open_date_str', 'N/A')
        
        header = doc.add_heading(level=1)
        run = header.add_run(f"{date_str} | {symbol} ({direction})")
        
        # 结果标记
        res_text = f"   {'✅ 盈利' if pnl > 0 else '❌ 亏损'} ${pnl:.2f}"
        res_run = header.add_run(res_text)
        if pnl > 0:
            res_run.font.color.rgb = RGBColor(0, 150, 0) # Green
        else:
            res_run.font.color.rgb = RGBColor(200, 0, 0) # Red
        
        # === 核心数据表格 (v7.0 增强版) ===
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '基础数据'
        hdr_cells[1].text = 'R倍数 / 波动率'
        hdr_cells[2].text = 'v7.0 心理/效率'
        
        # 数据行
        row_cells = table.add_row().cells
        
        # Col 1: 基础
        price_in = row.get('price', 0)
        # 尝试获取平仓价，如果没有则不显示
        row_cells[0].text = (
            f"策略: {row.get('strategy', '-')}\n"
            f"心态: {row.get('mental_state', '-')}\n"
            f"持续: {row.get('duration_str', '-')}"
        )
        
        # Col 2: R倍数 (MAE/MFE)
        mae = row.get('mae', '-')
        mfe = row.get('mfe', '-')
        mae_atr = row.get('mae_atr', None)
        
        mae_text = f"MAE: {mae} R"
        if mae_atr is not None and str(mae_atr) != 'nan':
            mae_text += f"\n({mae_atr:.1f}x ATR)" # 显示 ATR 倍数
            
        row_cells[1].text = (
            f"{mae_text}\n"
            f"MFE: {mfe} R\n"
            f"ETD: {row.get('etd', '-')} R"
        )
        
        # Col 3: v7.0 心理指标
        mad = row.get('mad', '-')
        eff = row.get('efficiency', '-')
        
        eff_str = f"{float(eff):.2f}" if (eff != '-' and str(eff) != 'nan') else "-"
        
        row_cells[2].text = (
            f"痛苦时长 (MAD): {mad} min\n"
            f"交易效率: {eff_str}\n"
            f"评分: {row.get('setup_rating', '-')}/10"
        )
        
        # === 交易笔记 (User Input) ===
        doc.add_heading('📝 你的复盘笔记:', level=3)
        notes = str(row.get('notes', '无笔记'))
        doc.add_paragraph(notes)
        
        # === 截图 (Image) ===
        # 支持截图字段，如果存在图片路径
        screenshot_path = row.get('screenshot', '')
        if screenshot_path and isinstance(screenshot_path, str):
            # 尝试多个可能的路径
            possible_paths = [
                screenshot_path,  # 直接路径
                os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads', screenshot_path),  # 相对路径
                os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'uploads', screenshot_path),  # Docker 路径
            ]
            
            img_found = False
            for img_path in possible_paths:
                if os.path.exists(img_path):
                    try:
                        doc.add_heading('📸 交易截图:', level=3)
                        doc.add_picture(img_path, width=Inches(5.0))
                        img_found = True
                        break
                    except Exception as e:
                        pass
            
            if not img_found and screenshot_path:
                # 如果所有路径都找不到，尝试从数据库路径推断
                # 这里假设 screenshot 是文件名，需要从数据库路径推断 uploads 目录
                pass
        
        # === AI 深度审计 (仅在 include_ai=True 时显示) ===
        if include_ai:
            ai_analysis = str(row.get('ai_analysis', ''))
            if ai_analysis and ai_analysis != 'None' and len(ai_analysis) > 5:
                doc.add_heading('🤖 AI 教练毒舌点评:', level=3)
                # 使用引用样式或斜体，区分 AI 内容
                p = doc.add_paragraph()
                runner = p.add_run(ai_analysis)
                runner.font.color.rgb = RGBColor(80, 80, 80) # 深灰色
                runner.italic = True
    
    # 保存文件
    doc.save(filename)
    return filename
